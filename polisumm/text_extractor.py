"""
Text Extraction Module
Handles extraction from PDF, DOCX, and TXT files
"""

import os
import re
from typing import Dict, List, Optional
import pdfplumber
try:
    from docx import Document
except ImportError:
    Document = None
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
from bs4 import BeautifulSoup
import html2text


class TextExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = True
        self.html_converter.ignore_images = True
        
    def extract(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, any]:
        """
        Extract text from a document
        
        Args:
            file_path: Path to the document
            file_type: Type of document (pdf, docx, txt). Auto-detected if None
            
        Returns:
            Dict with 'text', 'metadata', and 'sections'
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Auto-detect file type
        if file_type is None:
            file_type = self._detect_file_type(file_path)
        
        extractors = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt,
        }
        
        if file_type not in extractors:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return extractors[file_type](file_path)
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext[1:] if ext else 'txt'
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, any]:
        """Extract text from PDF using pdfplumber and PyMuPDF"""
        text_parts = []
        metadata = {
            'title': '',
            'author': '',
            'date': '',
            'pages': 0
        }
        
        try:
            # Try pdfplumber first for better text extraction
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            full_text = '\n\n'.join(text_parts)
            
            # Extract metadata using PyMuPDF if available
            if fitz is not None:
                doc = fitz.open(file_path)
                metadata_info = doc.metadata
                metadata.update({
                    'title': metadata_info.get('title', ''),
                    'author': metadata_info.get('author', ''),
                    'date': metadata_info.get('creationDate', ''),
                })
                doc.close()
            
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
        
        return {
            'text': full_text,
            'metadata': metadata,
            'sections': self._identify_sections(full_text)
        }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, any]:
        """Extract text from DOCX file"""
        if Document is None:
            raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            full_text = '\n'.join(paragraphs)
            
            metadata = {
                'title': '',
                'author': '',
                'date': '',
                'pages': len(doc.paragraphs)
            }
            
            # Extract core properties if available
            core_props = doc.core_properties
            if core_props:
                metadata['title'] = core_props.title or ''
                metadata['author'] = core_props.author or ''
                metadata['date'] = str(core_props.created or '')
            
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
        
        return {
            'text': full_text,
            'metadata': metadata,
            'sections': self._identify_sections(full_text)
        }
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()
        except Exception as e:
            raise Exception(f"Error extracting TXT: {str(e)}")
        
        return {
            'text': full_text,
            'metadata': {'title': os.path.basename(file_path), 'pages': 1},
            'sections': self._identify_sections(full_text)
        }
    
    def _identify_sections(self, text: str) -> List[Dict[str, any]]:
        """Identify document sections using headers and patterns"""
        sections = []
        lines = text.split('\n')
        
        current_section = {'title': 'Introduction', 'content': '', 'start_line': 0}
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Identify section headers (all caps, numbered, or prominent)
            if self._is_section_header(line_stripped):
                if current_section['content'].strip():
                    sections.append(current_section)
                
                current_section = {
                    'title': line_stripped,
                    'content': '',
                    'start_line': i
                }
            else:
                current_section['content'] += line + '\n'
        
        if current_section['content'].strip():
            sections.append(current_section)
        
        return sections
    
    def _is_section_header(self, line: str) -> bool:
        """Determine if a line is a section header"""
        if not line or len(line) < 3:
            return False
        
        # Check for patterns like "1. Title", "SECTION 1:", "Article 1"
        header_patterns = [
            r'^\d+[\.\)]',  # Numbered
            r'^SECTION\s+\d+',
            r'^Article\s+\d+',
            r'^CHAPTER\s+\d+',
            r'^[A-Z\s]{2,}$',  # All caps (short lines)
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        
        # Check if line is short and bold-looking
        if len(line) < 50 and line.isupper():
            return True
        
        return False
    
    def extract_from_html(self, html_content: str) -> Dict[str, any]:
        """Extract text from HTML content"""
        # Convert HTML to markdown first
        markdown_text = self.html_converter.handle(html_content)
        
        # Also extract pure text
        soup = BeautifulSoup(html_content, 'html.parser')
        plain_text = soup.get_text()
        
        return {
            'text': plain_text,
            'metadata': {'title': '', 'pages': 1},
            'sections': self._identify_sections(plain_text)
        }
    
    def extract_from_text(self, text: str) -> Dict[str, any]:
        """Handle direct text input"""
        return {
            'text': text,
            'metadata': {'title': 'Input Text', 'pages': 1},
            'sections': self._identify_sections(text)
        }

